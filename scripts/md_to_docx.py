"""Turn a paper draft in Markdown into the .docx a journal will accept.

Paper 1 was submitted as an editable Word file with the Korean version kept
alongside as reference, so the follow-ups keep that shape: English is the
submission, Korean is the draft.

Two things this deliberately does not do, both learned from paper 1's proofs:

- No field functions. Springer forbids them, so figure and table numbers are
  typed text rather than Word's Insert Caption.
- No bold inside tables unless the caller asks. Bold in a table triggers an
  author query asking what it signifies.

Usage:
    python scripts/md_to_docx.py paper3/draft_en.md paper3/draft_en.docx
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
FIG_RE = re.compile(r"\(`?(?:figures/)?(Fig\d+)\.png`?\)")


def add_runs(par, text: str) -> None:
    """Render **bold**, *italic* and `code` inside a paragraph."""
    for token in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            par.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            par.add_run(token[1:-1]).italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = par.add_run(token[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(9.5)
        else:
            par.add_run(token)


def add_table(doc: Document, rows: list[str]) -> None:
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], cells[2:]          # row 1 is the --- separator
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, name in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        add_runs(cell.paragraphs[0], name)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cs = table.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            cs[i].text = ""
            add_runs(cs[i].paragraphs[0], val)
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    lines = md_path.read_text(encoding="utf-8").split("\n")
    fig_dir = md_path.parent / "figures"

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line or line.strip() == "---":
            i += 1
            continue

        if line.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            par = doc.add_paragraph()
            run = par.add_run("\n".join(block))
            run.font.name = MONO_FONT
            run.font.size = Pt(9)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].replace("|", "").strip()) <= set("-: "):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append(lines[i])
                i += 1
            add_table(doc, rows)
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            doc.add_heading(line.lstrip("# ").strip(), level=min(level, 4))
            i += 1
            continue

        if line.startswith(("- ", "* ")):
            par = doc.add_paragraph(style="List Bullet")
            add_runs(par, line[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            par = doc.add_paragraph(style="List Number")
            add_runs(par, re.sub(r"^\d+\.\s", "", line))
            i += 1
            continue

        # A caption naming a figure file: place the image, then the caption.
        hit = FIG_RE.search(line)
        if hit:
            img = fig_dir / f"{hit.group(1)}.png"
            if img.exists():
                doc.add_picture(str(img), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            add_runs(cap, FIG_RE.sub("", line).strip())
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        par = doc.add_paragraph()
        add_runs(par, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"wrote {out_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: md_to_docx.py <input.md> <output.docx>")
    convert(Path(sys.argv[1]), Path(sys.argv[2]))
